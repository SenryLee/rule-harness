"""
Generate sample data files for the Rule Extraction Harness system.

Produces:
  - samples/采购合同模板.docx   Chinese procurement contract
  - samples/审查清单.xlsx       Excel review checklist (passthrough mode)
  - samples/案例.txt            Simulated court case excerpts

Usage:
  python samples/generate_samples.py
"""
from __future__ import annotations

from pathlib import Path

_SAMPLES_DIR = Path(__file__).resolve().parent


def _create_procurement_contract():
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("python-docx not installed; skipping 采购合同模板.docx")
        return

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_heading("采购合同", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contract_no = doc.add_paragraph("合同编号：PUR-2026-001")
    contract_no.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    doc.add_paragraph("买方：某科技有限公司")
    doc.add_paragraph("卖方：某制造有限公司")
    doc.add_paragraph("签订日期：2026年 月 日")
    doc.add_paragraph("签订地点：北京市")

    doc.add_paragraph("")

    articles = [
        ("第一条 合同标的与价格", [
            "1.1 买方同意向卖方购买，卖方同意向买方出售下列产品：",
            "   产品名称：工业级传感器模块",
            "   规格型号：ST-200 Pro",
            "   数量：10,000 个",
            "   单价：人民币 200 元/个（含税）",
            "   总价：人民币 2,000,000 元（大写：贰佰万元整）",
            "1.2 本合同总价为含增值税价格，税率为 13%。",
            "1.3 产品应符合附件一《技术规格书》的要求。",
        ]),
        ("第二条 付款条款", [
            "2.1 买方应在收到卖方开具的增值税专用发票后 30 日内支付合同总价的 90%，即人民币 1,800,000 元。",
            "2.2 剩余 10% 作为质保金，即人民币 200,000 元，在质保期满后 15 日内支付。",
            "2.3 质保期为产品验收合格之日起 12 个月。",
            "2.4 付款方式为银行转账。",
        ]),
        ("第三条 交付与验收", [
            "3.1 交货时间：卖方应于本合同生效后 60 日内完成全部产品的交付。",
            "3.2 交货地点：买方指定仓库（北京市海淀区）。",
            "3.3 验收标准：按照附件一《技术规格书》进行验收。",
            "3.4 验收期限：买方应在收到产品后 10 个工作日内完成验收。",
            "3.5 买方有权对不合格产品拒收，卖方应在 15 日内无条件更换。",
        ]),
        ("第四条 违约金", [
            "4.1 卖方逾期交付的，每日按迟延交付货物价值的 0.5% 支付违约金。",
            "4.2 累计违约金总额不得超过合同总价的 10%。",
            "4.3 买方逾期付款的，每日按逾期付款金额的 0.05% 支付违约金。",
            "4.4 违约金不足以弥补守约方损失的，违约方应赔偿差额部分。",
        ]),
        ("第五条 保密条款", [
            "5.1 双方应对在合同履行过程中知悉的对方商业秘密承担保密义务。",
            "5.2 商业秘密包括但不限于技术信息、客户信息、财务数据和经营策略。",
            "5.3 保密期限自合同终止之日起 5 年。",
            "5.4 保密信息不得用于本协议目的之外的任何用途。",
            "5.5 保密义务不适用于已公开信息、依法必须披露的信息。",
        ]),
        ("第六条 知识产权", [
            "6.1 因履行本合同产生的知识产权归买方所有。",
            "6.2 卖方保留其在本合同签订前已拥有的背景知识产权。",
            "6.3 卖方保证其提供的产品不侵犯任何第三方的知识产权。",
            "6.4 如发生知识产权侵权索赔，卖方应承担全部责任并赔偿买方损失。",
        ]),
        ("第七条 赔偿与责任", [
            "7.1 卖方对买方因卖方违约遭受的直接损失承担赔偿责任，赔偿总额不超过合同总价。",
            "7.2 卖方不对间接损失、利润损失或商誉损失承担责任。",
            "7.3 因卖方提供的产品存在缺陷导致第三方人身伤亡或财产损失的，卖方应承担赔偿责任，不受 7.1 条限制。",
        ]),
        ("第八条 争议解决", [
            "8.1 因本合同引起的争议，双方应协商解决。",
            "8.2 协商不成的，任何一方均有权提交北京仲裁委员会仲裁。",
            "8.3 仲裁裁决为终局的，对双方均有约束力。",
        ]),
        ("第九条 合同变更", [
            "9.1 本合同任何一方变更联系方式应提前 7 日书面通知对方。",
            "9.2 本合同的任何修改须经双方书面同意。",
        ]),
        ("第十条 法律适用", [
            "10.1 本合同的订立、效力、解释、履行及争议解决均适用中华人民共和国法律。",
        ]),
        ("第十一条 生效条款", [
            "11.1 本合同自双方法定代表人或授权代表签字并加盖公章之日起生效。",
            "11.2 本合同一式四份，双方各执两份，具有同等法律效力。",
        ]),
    ]

    for title_text, clauses in articles:
        doc.add_heading(title_text, level=1)
        for clause in clauses:
            p = doc.add_paragraph(clause)
            p.paragraph_format.first_line_indent = Cm(0)

    ending = [
        "",
        "买方（盖章）：                    卖方（盖章）：",
        "",
        "法定代表人或授权代表签字：         法定代表人或授权代表签字：",
        "",
        "日期：                            日期：",
    ]
    for line in ending:
        doc.add_paragraph(line)

    path = _SAMPLES_DIR / "采购合同模板.docx"
    doc.save(str(path))
    print(f"  Created: {path}")


def _create_review_checklist():
    try:
        import openpyxl
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    except ImportError:
        print("openpyxl not installed; skipping 审查清单.xlsx")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "审查清单"

    header_font = Font(name="SimHei", size=12, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    headers = ["序号", "风险点", "风险等级", "检查内容", "审查标准", "备注说明"]
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = thin_border

    rows: list[tuple[int, str, str, str, str, str]] = [
        (1, "违约金条款审查", "高", "是否约定违约金上限", "[条款] 违约金上限不超过合同总价30%", "参考民法典第585条，违约金过高可请求减少"),
        (2, "违约金计算基数", "高", "违约金计算基数是否明确", "[条款] 违约金以迟延交付货物价值为基数", ""),
        (3, "保密期限审查", "高", "保密期限是否明确", "[条款] 保密期限不少于3年", "核查保密期限起算点和终止条件"),
        (4, "保密范围界定", "中", "保密信息范围是否定义", "[条款] 商业秘密定义应包含技术、客户、财务信息", ""),
        (5, "知识产权归属", "高", "成果知识产权归属是否明确", "[条款] 履行合同产生的知识产权归买方所有", "注意区分前景与背景知识产权"),
        (6, "知识产权侵权担保", "高", "是否有第三方侵权担保条款", "[条款] 卖方保证产品不侵犯第三方知识产权", "应约定侵权赔偿机制"),
        (7, "付款条件审查", "中", "付款条件是否清晰可执行", "[条款] 收到发票后30日内支付", "核实发票类型及合规性"),
        (8, "质保金条款", "中", "质保金比例及释放条件", "[条款] 质保金为合同总价10%，质保期满后15日支付", ""),
        (9, "交付期限审查", "中", "交付时间是否明确", "[条款] 合同生效后60日内交付", "注意不可抗力等免责情形"),
        (10, "验收条款审查", "高", "验收标准和期限是否明确", "[条款] 按技术规格书验收，10个工作日内完成", "验收标准应客观可测量"),
        (11, "赔偿责任上限", "高", "赔偿责任上限是否合理", "[条款] 赔偿责任不超过合同总价", "注意排除故意或重大过失的责任上限"),
        (12, "争议解决方式", "中", "争议解决方式是否明确", "[条款] 协商不成提交仲裁", "注意仲裁机构名称及仲裁规则"),
        (13, "法律适用条款", "低", "法律适用是否约定", "[条款] 适用中华人民共和国法律", ""),
        (14, "合同生效条件", "低", "生效条件是否完备", "[条款] 双方签字并加盖公章后生效", "核实签约主体资格及授权"),
        (15, "通知送达条款", "低", "联系方式变更通知期限", "[条款] 变更联系方式提前7日通知", "建议增加送达地址确认条款"),
    ]

    body_font = Font(name="SimSun", size=11)
    body_align = Alignment(vertical="center", wrap_text=True)

    for row_idx, (idx, risk_point, risk_level, check_item, standard, notes) in enumerate(rows, 2):
        values = [idx, risk_point, risk_level, check_item, standard, notes]
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.font = body_font
            cell.alignment = body_align
            cell.border = thin_border

        risk_cell = ws.cell(row=row_idx, column=3)
        if risk_level == "高":
            risk_cell.font = Font(name="SimSun", size=11, bold=True, color="FF0000")
        elif risk_level == "中":
            risk_cell.font = Font(name="SimSun", size=11, bold=True, color="ED7D31")

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 20
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 28
    ws.column_dimensions["E"].width = 40
    ws.column_dimensions["F"].width = 30

    ws.auto_filter.ref = f"A1:F{len(rows) + 1}"

    path = _SAMPLES_DIR / "审查清单.xlsx"
    wb.save(str(path))
    print(f"  Created: {path}")


def _create_case_text():
    content = """\
某买卖合同纠纷案
====================================

案号：(2023)沪01民终12345号
审理法院：上海市第一中级人民法院

【基本案情】

原告（买方）与被告（卖方）于2022年签订《产品采购合同》，约定被告向原告供应
电子元器件10,000件，总价款人民币500万元。

合同履行过程中，原告发现部分产品存在质量问题，遂与被告协商退货。被告援引合同
第15.3条"最终解释权归本公司所有"拒绝退货退款请求。

此外，合同第15.4条载明"商品一经售出概不退换"。

原告遂诉至法院，请求判令：1. 确认上述条款无效；2. 被告退还货款并赔偿损失。

【法院判决】

法院经审理认为：

一、关于"最终解释权归本公司所有"条款

本案争议条款系出卖人为重复使用预先拟定的格式条款。该条款赋予出卖方单方解释合同
的权利，排除了相对方对合同条款的解释权，实质上限制并排除相对方的主要权利。

根据《中华人民共和国民法典》第四百九十七条规定："有下列情形之一的，该格式条款
无效：（一）具有本法第一编第六章第三节和本法第五百零六条规定的无效情形；
（二）提供格式条款一方不合理地免除或者减轻其责任、加重对方责任、限制对方主要
权利；（三）提供格式条款一方排除对方主要权利。"

案涉"最终解释权"条款属于排除对方主要权利的格式条款，应认定为无效。

二、关于"概不退换"条款

"概不退换"条款免除了出卖方对产品质量的担保义务，不合理地免除或减轻己方责任，
加重对方责任，违反了《民法典》第四百九十七条第（二）项之规定，亦属无效条款。

三、判决结果

1. 确认合同第15.3条"最终解释权归本公司所有"及第15.4条"商品一经售出概不退换"
   为无效格式条款；
2. 被告退还原告货款人民币500万元；
3. 驳回原告其他诉讼请求。


====================================
相关案例：北京某科技有限公司诉深圳某实业有限公司买卖合同纠纷案
====================================

案号：(2022)京0105民初67890号

【基本案情】

双方签订《软件销售合同》，约定买方逾期付款按日万分之五支付违约金。后买方逾期付款
60日，卖方主张违约金合计人民币9万元（相当于合同总价的18%）。

买方抗辩称违约金过高，请求法院予以调减。

【法院判决】

根据《民法典》第五百八十五条第二款："约定的违约金过分高于造成的损失的，人民法院
或者仲裁机构可以根据当事人的请求予以适当减少。"

法院结合守约方实际损失（主要为资金占用利息），认定约定的违约金过分高于实际损失，
将违约金调减至人民币3万元（约为年化利率6%计算的资金占用成本）。

法院同时指出：违约金条款的约定应兼顾惩罚违约行为和补偿守约方损失的双重功能，对
于违约金明显过高的条款，合同的相对方可在订立合同时即提出异议，避免履行阶段的争议。

本案提示：合同中违约金条款的设置应具有合理性，违约金比例过高的条款存在被司法机关
调减的风险。
"""
    path = _SAMPLES_DIR / "案例.txt"
    path.write_text(content, encoding="utf-8")
    print(f"  Created: {path}")


def main():
    _SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    print("Generating sample files...")
    _create_procurement_contract()
    _create_review_checklist()
    _create_case_text()
    print("Done.")


if __name__ == "__main__":
    main()
