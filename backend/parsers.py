from __future__ import annotations

import hashlib
import re
import shutil
import subprocess
from dataclasses import dataclass, field, replace
from pathlib import Path
from typing import Optional
from zipfile import BadZipFile


@dataclass(frozen=True)
class ContentBlock:
    block_id: str
    text: str
    location: str
    block_type: str


@dataclass(frozen=True)
class CommentBlock:
    comment_id: str
    author: str
    text: str
    anchor_location: str
    anchor_text: str


@dataclass(frozen=True)
class RevisionBlock:
    rev_id: str
    original_text: str
    revised_text: str
    location: str


@dataclass(frozen=True)
class ParsedDocument:
    sha256: str
    filename: str
    source_tag: str
    priority: int
    contract_types: list[str]
    industry_context: dict | None
    is_scanned: bool
    blocks: tuple[ContentBlock, ...]
    comments: tuple[CommentBlock, ...]
    revisions: tuple[RevisionBlock, ...]
    is_redline_doc: bool
    is_case_doc: bool
    is_passthrough: bool
    parse_warnings: tuple[str, ...] = ()


@dataclass(frozen=True)
class RuleCandidate:
    risk_level: str
    keywords: tuple[str, ...]
    check_item: str
    requirement: str
    notes: str
    rule_type: str
    theme_key: str
    subject: str
    predicate: str
    threshold_type: str
    direction: str
    source_excerpt: str
    source_location: str
    pipeline: str
    self_confidence: float
    uncertainty_points: tuple[str, ...]
    source_filename: str = ""
    source_sha256: str = ""
    source_tag: str = ""
    priority: int = 5
    contract_types: tuple[str, ...] = ()
    model: str = ""
    struct_check_pass: bool = True
    struct_failures: tuple[str, ...] = ()
    ladder: dict | None = None
    cited_cases: tuple[str, ...] | None = None
    combined_confidence: float = 0.0
    conflict_flag: str = "无"
    variant_versions: str = ""
    fingerprint: str = ""
    rule_id: str = ""
    enabled: str = "启用"
    jurisdiction: str = "中国大陆"
    # v1.1: 第五重门（忠实度）
    fidelity_pass: bool = True
    fidelity_failures: tuple[str, ...] = ()
    voice_match: bool = True
    output_target: str = "main"  # main / placeholder / negotiation / discarded
    task_mode: str = "full_library"
    scope_match: str = "in_scope"
    scope_reason: str = ""
    template_anchor: str = ""
    # 可审计分析摘要字段：进入 metadata/API，不进入主 CSV。
    assumption: str = ""
    behavior_mode: str = ""
    consequence: str = ""
    exception_conditions: str = ""
    review_action: str = ""
    transformation_note: str = ""


_DOCX_EXT = {".docx"}
_DOC_EXT = {".doc"}
_PDF_EXT = {".pdf"}
_XLSX_EXT = {".xlsx", ".xlsm", ".xls", ".csv", ".tsv"}
_TXT_EXT = {".txt", ".md", ".text"}
_SUPPORTED_EXT = _DOCX_EXT | _DOC_EXT | _PDF_EXT | _XLSX_EXT | _TXT_EXT

_PASSTHROUGH_HEADER_KEYWORDS = frozenset(
    {"风险", "检查", "审查", "标准", "说明", "要求", "关键词", "风险点", "风险程度",
     "风险等级", "风险分类", "控制措施", "检查项", "审查要点", "审核要点", "合规要求",
     "合规点", "风险描述", "审查内容", "审查依据", "管理要求", "具体要求"}
)

_SOURCE_PRIORITY_MAP = {
    "法规": 1,
    "监管文件": 1,
    "公司红线": 2,
    "谈判底线": 2,
    "内部制度": 3,
    "审批规则": 3,
    "标准条款库": 4,
    "示范文本": 4,
    "历史合同": 5,
    "合同模板": 5,
    "案例": 5,
    "争议材料": 5,
    "业务规范": 3,
    "行业特殊": 3,
    "审查清单": 3,
}

_HEADING_RX = re.compile(
    r"^[\s]*(?:第[一二三四五六七八九十百千0-9]+[章节条]"
    r"|[（(]?\s*[一二三四五六七八九十]+[、．.)）]"
    r"|§\d+|Article\s+\d+|"
    r"[0-9]+[.、．)]\s*)"
)

_RISK_LABEL_CN_MAP = {
    "高": {"name": "高"},
    "较高": {"name": "高"},
    "重大": {"name": "高"},
    "严重": {"name": "高"},
    "红": {"name": "高"},
    "中": {"name": "中"},
    "中等": {"name": "中"},
    "一般": {"name": "中"},
    "低": {"name": "低"},
    "较低": {"name": "低"},
    "轻微": {"name": "低"},
    "绿": {"name": "低"},
}

_COLUMN_KEYWORD_MAP: list[tuple[frozenset[str], str]] = [
    (frozenset({"关键词", "关键字", "触发词", "关键词汇", "keywords"}), "keywords"),
    (frozenset({"检查项", "检查内容", "检查要点", "检查", "check_item"}), "check_item"),
    (frozenset({"审查要求", "要求", "审查标准", "标准要求", "合规要求", "审查", "requirement"}), "requirement"),
    (frozenset({"审查说明", "说明", "备注", "审查建议", "建议", "补充说明", "notes"}), "notes"),
    (frozenset({"风险程度", "风险等级", "风险", "风险分类", "risk_level"}), "risk_level"),
    (frozenset({"适用合同类型", "合同类型", "适用类型", "contract_type"}), "contract_type"),
]


def _map_column_by_keywords(header: str, next_best: list[str]) -> str | None:
    header_norm = str(header).strip().lower()
    for candidates, field_name in _COLUMN_KEYWORD_MAP:
        for kw in candidates:
            if kw in header_norm:
                return field_name
    return None


def _resolve_source_priority(source_tag: str) -> int:
    return _SOURCE_PRIORITY_MAP.get(source_tag, 5)


def resolve_source_priority(source_tag: str) -> int:
    """Public wrapper used by the orchestrator."""
    return _resolve_source_priority(source_tag)


SOURCE_PRIORITY_MAP = _SOURCE_PRIORITY_MAP  # public alias


def compute_sha256(filepath: Path) -> str:
    hasher = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            hasher.update(chunk)
    return hasher.hexdigest()


def _make_block(block_id: str, text: str, location: str, block_type: str) -> ContentBlock:
    return ContentBlock(
        block_id=block_id,
        text=text.strip(),
        location=location,
        block_type=block_type,
    )


def parse_file(
    filepath: Path,
    source_tag: str,
    contract_types: list[str],
    industry_context: dict | None = None,
    is_scanned: bool = False,
    is_redline: bool = False,
    is_case: bool = False,
    ocr_enabled: bool = False,
    ocr_engine: str = "paddleocr",
    ocr_language: str = "ch+en",
) -> ParsedDocument:
    suffix = filepath.suffix.lower()
    sha256 = compute_sha256(filepath)
    priority = _resolve_source_priority(source_tag)

    if not filepath.exists():
        return ParsedDocument(
            sha256=sha256,
            filename=filepath.name,
            source_tag=source_tag,
            priority=priority,
            contract_types=list(contract_types),
            industry_context=industry_context,
            is_scanned=is_scanned,
            blocks=(),
            comments=(),
            revisions=(),
            is_redline_doc=is_redline,
            is_case_doc=is_case,
            is_passthrough=False,
        )

    if suffix in _DOCX_EXT:
        return parse_docx(filepath, source_tag, contract_types, industry_context,
                          is_scanned, is_redline, is_case, sha256, priority)
    elif suffix in _DOC_EXT:
        return parse_doc(filepath, source_tag, contract_types, industry_context,
                         is_scanned, is_redline, is_case, sha256, priority)
    elif suffix in _PDF_EXT:
        return parse_pdf(filepath, source_tag, contract_types, industry_context,
                         is_scanned, is_redline, is_case, sha256, priority,
                         ocr_enabled=ocr_enabled, ocr_engine=ocr_engine,
                         ocr_language=ocr_language)
    elif suffix in _XLSX_EXT:
        return parse_excel(filepath, source_tag, contract_types, industry_context,
                           is_scanned, is_redline, is_case, sha256, priority)
    elif suffix in _TXT_EXT:
        return parse_txt(filepath, source_tag, contract_types, industry_context,
                         is_scanned, is_redline, is_case, sha256, priority)
    else:
        return ParsedDocument(
            sha256=sha256,
            filename=filepath.name,
            source_tag=source_tag,
            priority=priority,
            contract_types=list(contract_types),
            industry_context=industry_context,
            is_scanned=is_scanned,
            blocks=(),
            comments=(),
            revisions=(),
            is_redline_doc=is_redline,
            is_case_doc=is_case,
            is_passthrough=False,
            parse_warnings=(f"unsupported_file_type:{suffix or '<none>'}",),
        )


def parse_doc(
    filepath: Path,
    source_tag: str,
    contract_types: list[str],
    industry_context: dict | None = None,
    is_scanned: bool = False,
    is_redline: bool = False,
    is_case: bool = False,
    sha256: str | None = None,
    priority: int | None = None,
) -> ParsedDocument:
    """Parse legacy Word .doc files through available system converters."""
    if sha256 is None:
        sha256 = compute_sha256(filepath)
    if priority is None:
        priority = _resolve_source_priority(source_tag)

    text, warning = _extract_doc_text(filepath)
    if not text.strip():
        return _empty_parsed(
            filepath.name,
            source_tag,
            contract_types,
            industry_context,
            is_scanned,
            is_redline,
            is_case,
            sha256,
            priority,
            parse_warnings=(warning or "legacy_doc_no_text",),
        )

    blocks = tuple(_chunk_by_double_newline(text, "doc", 0))
    if not blocks:
        blocks = (
            ContentBlock(
                block_id="doc-0",
                text=text.strip(),
                location="doc-0",
                block_type="paragraph",
            ),
        )

    return ParsedDocument(
        sha256=sha256,
        filename=filepath.name,
        source_tag=source_tag,
        priority=priority,
        contract_types=list(contract_types),
        industry_context=industry_context,
        is_scanned=is_scanned,
        blocks=blocks,
        comments=(),
        revisions=(),
        is_redline_doc=is_redline,
        is_case_doc=is_case,
        is_passthrough=False,
        parse_warnings=(),
    )


def _extract_doc_text(filepath: Path) -> tuple[str, str | None]:
    commands = [
        ("textutil", ["textutil", "-convert", "txt", "-stdout", str(filepath)]),
        ("antiword", ["antiword", str(filepath)]),
        ("catdoc", ["catdoc", str(filepath)]),
    ]
    missing: list[str] = []
    failures: list[str] = []

    for tool, cmd in commands:
        if shutil.which(tool) is None:
            missing.append(tool)
            continue
        try:
            proc = subprocess.run(
                cmd,
                check=False,
                capture_output=True,
                text=True,
                timeout=45,
            )
        except Exception as exc:
            failures.append(f"{tool}:{type(exc).__name__}")
            continue
        if proc.returncode == 0 and proc.stdout.strip():
            return proc.stdout, None
        detail = (proc.stderr or proc.stdout or "").strip().replace("\n", " ")
        failures.append(f"{tool}:{proc.returncode}:{detail[:120]}")

    if failures:
        return "", "legacy_doc_conversion_failed:" + "|".join(failures)
    return "", "legacy_doc_converter_missing:" + ",".join(missing)


def parse_docx(
    filepath: Path,
    source_tag: str,
    contract_types: list[str],
    industry_context: dict | None = None,
    is_scanned: bool = False,
    is_redline: bool = False,
    is_case: bool = False,
    sha256: str | None = None,
    priority: int | None = None,
) -> ParsedDocument:
    if sha256 is None:
        sha256 = compute_sha256(filepath)
    if priority is None:
        priority = _resolve_source_priority(source_tag)

    try:
        import docx
    except ImportError:
        return _empty_parsed(filepath.name, source_tag, contract_types,
                             industry_context, is_scanned, is_redline, is_case,
                             sha256, priority)

    try:
        document = docx.Document(str(filepath))
    except (BadZipFile, ValueError, OSError):
        return _empty_parsed(filepath.name, source_tag, contract_types,
                             industry_context, is_scanned, is_redline, is_case,
                             sha256, priority)

    blocks: list[ContentBlock] = []
    comments: list[CommentBlock] = []
    revisions: list[RevisionBlock] = []
    is_passthrough = False

    para_idx = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = str(paragraph.style.name) if paragraph.style else ""
        block_type = "paragraph"
        location = str(para_idx)

        if "Heading" in style_name or _HEADING_RX.match(text):
            block_type = "heading"

        blocks.append(_make_block(
            block_id=f"p{para_idx}",
            text=text,
            location=location,
            block_type=block_type,
        ))
        para_idx += 1

    table_blocks, table_is_passthrough = _extract_docx_table_blocks(document)
    blocks.extend(table_blocks)
    is_passthrough = table_is_passthrough

    _extract_docx_comments(filepath, comments)
    _extract_docx_revisions(document, revisions)

    return ParsedDocument(
        sha256=sha256,
        filename=filepath.name,
        source_tag=source_tag,
        priority=priority,
        contract_types=list(contract_types),
        industry_context=industry_context,
        is_scanned=is_scanned,
        blocks=tuple(blocks),
        comments=tuple(comments),
        revisions=tuple(revisions),
        is_redline_doc=is_redline,
        is_case_doc=is_case,
        is_passthrough=is_passthrough,
    )


def _extract_docx_table_blocks(document) -> tuple[list[ContentBlock], bool]:
    blocks: list[ContentBlock] = []
    is_passthrough = False

    for table_idx, table in enumerate(document.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue

        headers = rows[0]
        table_passthrough = _compute_passthrough_score(headers) >= 3
        is_passthrough = is_passthrough or table_passthrough
        data_rows = rows[1:] if table_passthrough else rows

        for row_idx, cells in enumerate(data_rows, start=2 if table_passthrough else 1):
            if table_passthrough:
                parts = []
                for col_idx, cell in enumerate(cells):
                    header = headers[col_idx] if col_idx < len(headers) else f"Col{col_idx + 1}"
                    if cell:
                        parts.append(f"{header}: {cell}")
                text = "; ".join(parts)
            else:
                text = " | ".join(cell for cell in cells if cell)
            if not text.strip():
                continue
            blocks.append(ContentBlock(
                block_id=f"t{table_idx + 1}r{row_idx}",
                text=text,
                location=f"table-{table_idx + 1}-row-{row_idx}",
                block_type="table_row",
            ))

    return blocks, is_passthrough


def _extract_docx_comments(filepath: Path, comments_out: list[CommentBlock]) -> None:
    try:
        from zipfile import ZipFile
        from lxml import etree
    except ImportError:
        return

    try:
        with ZipFile(str(filepath), "r") as zf:
            if "word/comments.xml" not in zf.namelist():
                return
            comments_xml = zf.read("word/comments.xml")
    except (BadZipFile, KeyError, OSError):
        return

    try:
        root = etree.fromstring(comments_xml)
    except etree.XMLSyntaxError:
        return

    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    namespaces = {"w": w_ns}

    for comment_elem in root.findall(f"{{{w_ns}}}comment"):
        comment_id = comment_elem.get(f"{{{w_ns}}}id", "")
        author = comment_elem.get(f"{{{w_ns}}}author", "")

        text_parts = []
        for t_elem in comment_elem.iter(f"{{{w_ns}}}t"):
            if t_elem.text:
                text_parts.append(t_elem.text)

        anchor = comment_elem.get(f"{{{w_ns}}}anchor", "")
        comment_text = "".join(text_parts).strip()

        if comment_text:
            comments_out.append(CommentBlock(
                comment_id=comment_id,
                author=author,
                text=comment_text,
                anchor_location=anchor,
                anchor_text="",
            ))


def _extract_docx_revisions(document, revisions_out: list[RevisionBlock]) -> None:
    """
    遍历每个段落，识别 w:ins / w:del 节点。每个含修订的段落输出 1 条 RevisionBlock：
      - revised_text: 段落当前可见文本（含 ins，去除 del）
      - original_text: 段落原文（含 del，去除 ins）
    """
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for para_idx, paragraph in enumerate(document.paragraphs):
        ins_segments: list[str] = []
        del_segments: list[str] = []
        has_any = False

        for child in paragraph._element.iter():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "ins":
                # 收集 ins 内所有 w:t 文本
                for t in child.iter(f"{{{w_ns}}}t"):
                    if t.text:
                        ins_segments.append(t.text)
                        has_any = True
            elif tag == "del":
                # w:del 里通常是 w:delText
                for dt in child.iter():
                    dtag = dt.tag.split("}")[-1] if "}" in dt.tag else dt.tag
                    if dtag in ("delText", "t") and dt.text:
                        del_segments.append(dt.text)
                        has_any = True

        if not has_any:
            continue

        current_visible = paragraph.text.strip()
        deleted_text = "".join(del_segments).strip()
        inserted_text = "".join(ins_segments).strip()

        original = (current_visible.replace(inserted_text, "") if inserted_text else current_visible)
        original = (original + deleted_text).strip()
        revised = current_visible

        revisions_out.append(RevisionBlock(
            rev_id=f"r{para_idx}",
            original_text=original or current_visible,
            revised_text=revised,
            location=f"paragraph-{para_idx}",
        ))


def parse_pdf(
    filepath: Path,
    source_tag: str,
    contract_types: list[str],
    industry_context: dict | None = None,
    is_scanned: bool = False,
    is_redline: bool = False,
    is_case: bool = False,
    sha256: str | None = None,
    priority: int | None = None,
    ocr_enabled: bool = False,
    ocr_engine: str = "paddleocr",
    ocr_language: str = "ch+en",
) -> ParsedDocument:
    if sha256 is None:
        sha256 = compute_sha256(filepath)
    if priority is None:
        priority = _resolve_source_priority(source_tag)

    blocks = _extract_pdf_text(filepath, is_scanned)
    scanned = is_scanned or _detect_scanned(blocks)
    warnings: list[str] = []
    if scanned and not blocks and ocr_enabled:
        blocks, ocr_warning = _ocr_pdf_pages(filepath, ocr_engine, ocr_language)
        if ocr_warning:
            warnings.append(ocr_warning)
    elif scanned and not blocks:
        warnings.append("pdf_scanned_without_ocr")

    return ParsedDocument(
        sha256=sha256,
        filename=filepath.name,
        source_tag=source_tag,
        priority=priority,
        contract_types=list(contract_types),
        industry_context=industry_context,
        is_scanned=scanned,
        blocks=blocks,
        comments=(),
        revisions=(),
        is_redline_doc=is_redline,
        is_case_doc=is_case,
        is_passthrough=False,
        parse_warnings=tuple(warnings),
    )


def _extract_pdf_text(filepath: Path, is_scanned: bool) -> tuple[ContentBlock, ...]:
    blocks: list[ContentBlock] = []
    block_counter = 0

    try:
        import pdfplumber

        with pdfplumber.open(str(filepath)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    page_blocks = _chunk_by_double_newline(
                        text, f"p{page_num + 1}", block_counter
                    )
                    blocks.extend(page_blocks)
                    block_counter += len(page_blocks)
        if blocks:
            return tuple(blocks)
    except Exception:
        pass

    try:
        import fitz

        doc = fitz.open(str(filepath))
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text:
                page_blocks = _chunk_by_double_newline(
                    text, f"p{page_num + 1}", block_counter
                )
                blocks.extend(page_blocks)
                block_counter += len(page_blocks)
        doc.close()
    except Exception:
        pass

    return tuple(blocks)


def _chunk_by_double_newline(
    text: str, page_label: str, start_counter: int
) -> list[ContentBlock]:
    chunks = re.split(r"\n\s*\n", text)
    result: list[ContentBlock] = []
    for i, chunk in enumerate(chunks):
        stripped = chunk.strip()
        if not stripped:
            continue
        result.append(ContentBlock(
            block_id=f"{page_label}-{i}",
            text=stripped,
            location=f"{page_label}-{i}",
            block_type="paragraph",
        ))
    return result


def _detect_scanned(blocks: tuple[ContentBlock, ...]) -> bool:
    if not blocks:
        return True
    total_chars = sum(len(b.text) for b in blocks)
    total_blocks = len(blocks)
    if total_blocks == 0:
        return True
    avg_chars = total_chars / total_blocks
    return avg_chars < 10 and total_chars < 100


def _ocr_pdf_pages(
    filepath: Path,
    engine: str,
    language: str,
) -> tuple[tuple[ContentBlock, ...], str | None]:
    """Best-effort OCR for ordinary scanned PDFs.

    The implementation intentionally treats OCR engines as optional runtime
    dependencies. If PaddleOCR or Tesseract is unavailable, parsing continues
    with a warning instead of crashing the batch.
    """
    blocks: list[ContentBlock] = []
    warning: str | None = None

    try:
        import fitz
    except Exception:
        return (), "pdf_ocr_unavailable:pymupdf_missing"

    try:
        doc = fitz.open(str(filepath))
    except Exception:
        return (), "pdf_ocr_unavailable:open_failed"

    try:
        for page_num in range(len(doc)):
            text, page_warning = _ocr_pdf_page(doc[page_num], engine, language)
            if page_warning and warning is None:
                warning = page_warning
            if not text.strip():
                continue
            for block in _chunk_by_double_newline(text, f"p{page_num + 1}-ocr", len(blocks)):
                blocks.append(block)
    finally:
        doc.close()

    if not blocks and warning is None:
        warning = "pdf_ocr_no_text"
    return tuple(blocks), warning


def _ocr_pdf_page(page, engine: str, language: str) -> tuple[str, str | None]:
    try:
        import tempfile
    except Exception:
        return "", "pdf_ocr_unavailable:tempfile_missing"

    matrix = getattr(page, "get_pixmap")
    try:
        pix = matrix(dpi=220, alpha=False)
    except TypeError:
        pix = matrix()

    suffix = ".png"
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=True) as tmp:
        pix.save(tmp.name)
        normalized_engine = (engine or "").lower()
        if normalized_engine == "paddleocr":
            text, warning = _ocr_image_with_paddle(tmp.name, language)
            if text or warning != "pdf_ocr_unavailable:paddleocr_missing":
                return text, warning
        return _ocr_image_with_tesseract(tmp.name, language)


def _ocr_image_with_paddle(image_path: str, language: str) -> tuple[str, str | None]:
    try:
        from paddleocr import PaddleOCR
    except Exception:
        return "", "pdf_ocr_unavailable:paddleocr_missing"

    lang = "ch" if "ch" in (language or "").lower() else "en"
    try:
        ocr = PaddleOCR(use_angle_cls=True, lang=lang, show_log=False)
        result = ocr.ocr(image_path, cls=True)
    except Exception as exc:
        return "", f"pdf_ocr_failed:paddleocr:{exc}"

    lines: list[str] = []
    for page_result in result or []:
        for item in page_result or []:
            try:
                text = item[1][0]
            except (IndexError, TypeError):
                continue
            if text:
                lines.append(str(text))
    return "\n".join(lines), None


def _ocr_image_with_tesseract(image_path: str, language: str) -> tuple[str, str | None]:
    import shutil
    import subprocess

    if shutil.which("tesseract") is None:
        return "", "pdf_ocr_unavailable:tesseract_missing"

    lang = _tesseract_lang(language)
    cmd = ["tesseract", image_path, "stdout"]
    if lang:
        cmd.extend(["-l", lang])
    try:
        proc = subprocess.run(cmd, check=False, capture_output=True, text=True, timeout=60)
    except Exception as exc:
        return "", f"pdf_ocr_failed:tesseract:{exc}"
    if proc.returncode != 0:
        return "", f"pdf_ocr_failed:tesseract:{proc.stderr.strip()[:120]}"
    return proc.stdout.strip(), None


def _tesseract_lang(language: str) -> str:
    raw = (language or "").lower()
    if raw in {"ch+en", "zh+en", "chi_sim+eng"}:
        return "chi_sim+eng"
    if raw in {"ch", "zh", "chi_sim"}:
        return "chi_sim"
    if raw in {"en", "eng"}:
        return "eng"
    return raw.replace("ch", "chi_sim").replace("en", "eng")


def parse_excel(
    filepath: Path,
    source_tag: str,
    contract_types: list[str],
    industry_context: dict | None = None,
    is_scanned: bool = False,
    is_redline: bool = False,
    is_case: bool = False,
    sha256: str | None = None,
    priority: int | None = None,
) -> ParsedDocument:
    if sha256 is None:
        sha256 = compute_sha256(filepath)
    if priority is None:
        priority = _resolve_source_priority(source_tag)

    try:
        import openpyxl
        from openpyxl.utils.exceptions import InvalidFileException
    except ImportError:
        return _empty_parsed(filepath.name, source_tag, contract_types,
                             industry_context, is_scanned, is_redline, is_case,
                             sha256, priority)

    suffix = filepath.suffix.lower()
    if suffix in (".csv", ".tsv"):
        return _parse_csv_as_passthrough(filepath, source_tag, contract_types,
                                         industry_context, is_scanned, is_redline,
                                         is_case, sha256, priority)

    if suffix == ".xls":
        # openpyxl 只读 OOXML，旧版 .xls（OLE2）读不了；明确提示而非静默空结果。
        return _empty_parsed(filepath.name, source_tag, contract_types,
                             industry_context, is_scanned, is_redline, is_case,
                             sha256, priority,
                             parse_warnings=("xls_unsupported_convert_to_xlsx",))

    try:
        wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
    except (BadZipFile, OSError, KeyError, InvalidFileException):
        return _empty_parsed(filepath.name, source_tag, contract_types,
                             industry_context, is_scanned, is_redline, is_case,
                             sha256, priority)

    blocks: list[ContentBlock] = []
    is_passthrough = False
    multi_sheet = len(wb.sheetnames) > 1

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(min_row=1, values_only=True))
        if not rows:
            continue
        sheet_blocks, sheet_passthrough = _sheet_rows_to_blocks(
            sheet_name, rows, multi_sheet
        )
        is_passthrough = is_passthrough or sheet_passthrough
        blocks.extend(sheet_blocks)

    wb.close()

    return ParsedDocument(
        sha256=sha256,
        filename=filepath.name,
        source_tag=source_tag,
        priority=priority,
        contract_types=list(contract_types),
        industry_context=industry_context,
        is_scanned=is_scanned,
        blocks=tuple(blocks),
        comments=(),
        revisions=(),
        is_redline_doc=is_redline,
        is_case_doc=is_case,
        is_passthrough=is_passthrough,
    )


def _parse_csv_as_passthrough(
    filepath: Path,
    source_tag: str,
    contract_types: list[str],
    industry_context: dict | None,
    is_scanned: bool,
    is_redline: bool,
    is_case: bool,
    sha256: str,
    priority: int,
) -> ParsedDocument:
    import csv

    blocks: list[ContentBlock] = []
    is_passthrough = False

    with open(filepath, "r", encoding="utf-8-sig", newline="") as f:
        delimiter = "\t" if filepath.suffix.lower() == ".tsv" else ","
        reader = csv.reader(f, delimiter=delimiter)
        all_rows = list(reader)

    if not all_rows:
        return _empty_parsed(filepath.name, source_tag, contract_types,
                             industry_context, is_scanned, is_redline, is_case,
                             sha256, priority)

    # 与 xlsx 共用同一套行→块逻辑：检测表头、按行打"表头: 值"标注，避免整表拼成一个巨块。
    sheet_blocks, is_passthrough = _sheet_rows_to_blocks("Sheet1", all_rows, multi_sheet=False)
    blocks.extend(sheet_blocks)

    return ParsedDocument(
        sha256=sha256,
        filename=filepath.name,
        source_tag=source_tag,
        priority=priority,
        contract_types=list(contract_types),
        industry_context=industry_context,
        is_scanned=is_scanned,
        blocks=tuple(blocks),
        comments=(),
        revisions=(),
        is_redline_doc=is_redline,
        is_case_doc=is_case,
        is_passthrough=is_passthrough,
    )


def _compute_passthrough_score(headers: list[str]) -> int:
    score = 0
    for h in headers:
        h_norm = str(h).strip().lower()
        for kw in _PASSTHROUGH_HEADER_KEYWORDS:
            if kw in h_norm:
                score += 1
                break
    return score


# 表格解析最多向下扫描多少行寻找真正的表头（跳过标题/横幅/空行）。
_HEADER_SCAN_ROWS = 8


def _col_letter(idx0: int) -> str:
    """0 基列号 → Excel 列字母（0→A, 26→AA）。用于给无表头的列兜底命名。"""
    s = ""
    n = idx0 + 1
    while n:
        n, r = divmod(n - 1, 26)
        s = chr(65 + r) + s
    return s


def _cell_str(cell: object) -> str:
    return str(cell).strip() if cell is not None else ""


def _detect_header_row(rows: list, scan: int = _HEADER_SCAN_ROWS) -> int | None:
    """在前若干行里挑"最像表头"的一行：先比直通关键词命中数，再比非空单元格数。

    返回行下标；若前若干行都不足 2 个非空单元格（多半是叙述型/单列表），返回 None。
    这样标题行、横幅行、空行在真正表头之上时不会被误当表头，避免列错位。
    """
    candidates: list[tuple[int, int]] = []  # (行下标, 关键词得分)
    for i, row in enumerate(rows[:scan]):
        nonempty = sum(1 for c in row if _cell_str(c))
        if nonempty < 2:
            continue
        score = _compute_passthrough_score([_cell_str(c) for c in row])
        candidates.append((i, score))
    if not candidates:
        return None
    # 仅当某行命中审查类关键词达到直通阈值(≥3)时，才认定它是规则库表头并跨过其上的标题行；
    # 阈值以下的零星命中（如数据里出现"标准间"含"标准"）不作数，避免把数据行误当表头。
    strong = [(i, s) for i, s in candidates if s >= 3]
    if strong:
        max_score = max(s for _, s in strong)
        return min(i for i, s in strong if s == max_score)
    # 无强关键词信号（普通内容表）：表头在数据之上，取最靠上的"≥2 非空"行，
    # 不能按非空单元格数挑，否则会把单元格更满的数据行误当表头。
    return candidates[0][0]


def _row_to_labeled_text(headers: list[str], row: tuple) -> str:
    """把一行渲染成 "表头: 值; 表头: 值"，跳过空单元格，空表头用 列X 兜底。

    保留列语义，避免整行单元格用空格拼接后串味/错位。
    """
    parts: list[str] = []
    for ci, cell in enumerate(row):
        val = _cell_str(cell)
        if not val:
            continue
        label = headers[ci] if ci < len(headers) and headers[ci] else f"列{_col_letter(ci)}"
        parts.append(f"{label}: {val}")
    return "; ".join(parts)


def _row_to_plain_text(row: tuple) -> str:
    """无表头时：把行内非空单元格用 ' | ' 连接，保留单元格边界。"""
    return " | ".join(_cell_str(c) for c in row if _cell_str(c))


def _sheet_rows_to_blocks(
    sheet_name: str,
    rows: list,
    multi_sheet: bool,
) -> tuple[list[ContentBlock], bool]:
    """把单个工作表/CSV 的行转成块，返回 (blocks, 该表是否为直通规则表)。

    - 检测表头行（跳过其上的标题/空行）；
    - 直通表（表头命中审查类关键词 ≥3）：每条数据行 → 一个 table_row 块，交直通管道按行导入；
    - 内容表：每条数据行 → 一个带"表头: 值"标注的 paragraph 块（不再整表拼成一个巨块），
      多表时在表前加一个工作表名 heading 块作为上下文；
    - 无表头（叙述/单列）：每非空行 → 一个 paragraph 块。
    """
    blocks: list[ContentBlock] = []
    header_idx = _detect_header_row(rows)

    if header_idx is None:
        for r, row in enumerate(rows, start=1):
            text = _row_to_plain_text(row)
            if text:
                blocks.append(ContentBlock(
                    block_id=f"{sheet_name}!A{r}",
                    text=text,
                    location=f"{sheet_name}!A{r}",
                    block_type="paragraph",
                ))
        return blocks, False

    headers = [_cell_str(c) for c in rows[header_idx]]
    passthrough = _compute_passthrough_score(headers) >= 3
    block_type = "table_row" if passthrough else "paragraph"

    if multi_sheet and not passthrough:
        blocks.append(ContentBlock(
            block_id=f"{sheet_name}!hdr",
            text=f"【工作表：{sheet_name}】",
            location=sheet_name,
            block_type="heading",
        ))

    for r, row in enumerate(rows[header_idx + 1:], start=header_idx + 2):
        text = _row_to_labeled_text(headers, row)
        if not text:
            continue
        blocks.append(ContentBlock(
            block_id=f"{sheet_name}!A{r}",
            text=text,
            location=f"{sheet_name}!A{r}",
            block_type=block_type,
        ))
    return blocks, passthrough


def _empty_parsed(
    filename: str,
    source_tag: str,
    contract_types: list[str],
    industry_context: dict | None,
    is_scanned: bool,
    is_redline: bool,
    is_case: bool,
    sha256: str,
    priority: int,
    parse_warnings: tuple[str, ...] = (),
) -> ParsedDocument:
    return ParsedDocument(
        sha256=sha256,
        filename=filename,
        source_tag=source_tag,
        priority=priority,
        contract_types=list(contract_types),
        industry_context=industry_context,
        is_scanned=is_scanned,
        blocks=(),
        comments=(),
        revisions=(),
        is_redline_doc=is_redline,
        is_case_doc=is_case,
        is_passthrough=False,
        parse_warnings=parse_warnings,
    )


def map_passthrough_row_to_fields(row_text: str, headers: list[str]) -> dict[str, str]:
    field_map: dict[str, str] = {}
    pairs = row_text.split("; ")
    for pair in pairs:
        if ": " not in pair:
            continue
        header, value = pair.split(": ", 1)
        header_norm = header.strip().lower()
        mapped_field = _map_column_by_keywords(header_norm, [])
        if mapped_field:
            field_map[mapped_field] = value.strip()
    return field_map


def normalize_risk_label(label: str) -> str:
    for key, entry in _RISK_LABEL_CN_MAP.items():
        if key in label:
            return entry["name"]
    return "中"


def parse_txt(
    filepath: Path,
    source_tag: str,
    contract_types: list[str],
    industry_context: dict | None = None,
    is_scanned: bool = False,
    is_redline: bool = False,
    is_case: bool = False,
    sha256: str | None = None,
    priority: int | None = None,
) -> ParsedDocument:
    if sha256 is None:
        sha256 = compute_sha256(filepath)
    if priority is None:
        priority = _resolve_source_priority(source_tag)

    raw_text = filepath.read_text(encoding="utf-8", errors="replace")
    paragraphs = [p.strip() for p in raw_text.split("\n\n") if p.strip()]

    blocks: list[ContentBlock] = []
    for i, para in enumerate(paragraphs):
        blocks.append(ContentBlock(
            block_id=f"txt-{i}",
            text=para,
            location=f"paragraph-{i+1}",
            block_type="paragraph",
        ))

    return ParsedDocument(
        sha256=sha256,
        filename=filepath.name,
        source_tag=source_tag,
        priority=priority,
        contract_types=list(contract_types),
        industry_context=industry_context,
        is_scanned=is_scanned,
        blocks=tuple(blocks),
        comments=(),
        revisions=(),
        is_redline_doc=is_redline,
        is_case_doc=is_case,
        is_passthrough=False,
    )
