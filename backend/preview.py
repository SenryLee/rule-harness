from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

import yaml

from backend.config import PROJECT_ROOT
from backend.document_profile import profile_document

PROFILES_DIR = PROJECT_ROOT / "profiles"

# v1.2：词表统一迁移到 backend/classification/taxonomy.yaml；下面是加载失败的兜底。
_FALLBACK_SOURCE_RULES: list[dict[str, Any]] = [
    {
        "label": "公司红线",
        "filename": ("红线", "底线", "退让", "不可接受", "谈判清单"),
        "body": ("红线", "底线", "不可接受"),
        "reason": "命中红线/底线类来源线索",
    },
    {
        "label": "标准条款库",
        "filename": ("手册", "指引", "指南", "规范", "审核要点", "审查清单", "合同审查", "规则项导入", "审查意见"),
        "body": ("审核要点", "审查清单", "规则项", "标准条款"),
        "reason": "命中审查手册/规则库类来源线索",
    },
    {
        "label": "法规",
        "filename": ("法律法规", "法律", "法规", "条例", "办法", "司法解释", "规定", "细则", "规程"),
        "body": ("法律法规", "法规", "条例", "司法解释", "应当遵守"),
        "reason": "命中法规/细则类来源线索",
    },
    {
        "label": "合同模板",
        "filename": ("模板", "范本", "示范文本", "合同样本"),
        "body": ("模板", "范本", "示范文本"),
        "reason": "命中模板/范本类来源线索",
    },
    {
        "label": "内部制度",
        "filename": ("制度", "流程", "内控", "操作规程", "管理规范"),
        "body": ("内部制度", "内控", "操作流程"),
        "reason": "命中制度/流程类来源线索",
    },
    {
        "label": "案例",
        "filename": ("案例", "判例", "判决", "裁判", "纠纷案", "抗诉案", "检例", "再审", "民事判决"),
        "body": ("案例", "判例", "法院认为", "裁判观点"),
        "reason": "命中案例/裁判类来源线索",
    },
]

_CASE_FILENAME_RX = re.compile(r"(?:^|[^\w])[^，。\n]{1,16}诉[^，。\n]{1,32}(?:纠纷|案)")


def _source_rules() -> list[dict[str, Any]]:
    from backend.classification import load_source_rules

    return load_source_rules() or _FALLBACK_SOURCE_RULES

_PROFILE_STRONG_KEYWORDS: dict[str, tuple[str, ...]] = {
    "建工·总包": (
        "EPC", "设计采购施工", "工程总承包", "施工总承包", "总承包", "发包人", "承包人",
        "工程量清单", "施工许可", "竣工验收", "竣工结算", "工期索赔", "工程变更",
        "质量保证金", "建设项目", "工程建设", "施工单位", "监理",
    ),
    "建工·勘察设计": ("勘察", "设计合同", "初步设计", "施工图设计", "设计成果", "设计变更"),
    "房地产": (
        "房地产", "房产", "房屋", "不动产", "商品房", "物业", "业主", "产权",
        "房屋赠与", "赠与房产", "房产赠与", "过户", "登记", "公证", "共有",
    ),
    "金融": (
        "银行", "证券", "保险", "资管", "信托", "融资", "保理", "基金", "金融",
        "反洗钱", "投资者适当性", "衍生品", "授信",
    ),
    "医药": ("药品", "医疗器械", "临床试验", "GMP", "GSP", "医保", "处方"),
    "IT": ("软件", "系统集成", "SaaS", "API", "源代码", "数据接口", "网络安全", "数据出境"),
    "制造": ("生产线", "设备", "模具", "来料加工", "质量检验", "交付验收", "质保"),
    "能源·电力": ("能源", "电力", "光伏", "风电", "储能", "并网", "购售电", "电站"),
    "汽车": ("汽车", "整车", "零部件", "主机厂", "经销商", "召回", "三包"),
    "通用商事": (
        "买卖", "销售", "采购", "服务", "代理", "经销", "合作", "赠与合同",
        "赠与", "合同无效", "解除条件", "违约责任", "争议解决", "动产",
    ),
}

_GENERIC_PROFILE_TERMS = {
    "合同", "条款", "法律", "适用法律", "协商", "调解", "诉讼", "争议解决", "管辖",
    "变更", "不可抗力", "服务", "保证", "担保", "数量", "验收", "质量标准",
}

_SOURCE_AUTO_THRESHOLD = 0.55
_CONTRACT_AUTO_THRESHOLD = 0.62
_PARTY_AUTO_THRESHOLD = 0.55

_PARTY_RULES: list[tuple[str, tuple[str, ...]]] = [
    ("发包人", ("发包人", "建设单位", "业主方")),
    ("承包人", ("承包人", "施工单位", "总承包人")),
    ("甲方", ("甲方",)),
    ("乙方", ("乙方",)),
    ("出租人", ("出租人",)),
    ("承租人", ("承租人",)),
    ("买方", ("买方", "采购方")),
    ("卖方", ("卖方", "供应商")),
]


def extract_preview_text(filename: str, content: bytes, limit: int = 3000) -> str:
    """Return preview text for classification（v1.2：头/中/尾多点采样）。

    判决书的"本院认为"、法规的附则、手册的章节正文都在中后部——
    只取头部 2048 字是旧版误判的最大来源。
    本函数从不调用 LLM；解析依赖缺失时退回字节前缀解码。
    """
    text, _ = extract_preview_sample(filename, content, limit)
    return text


def extract_preview_sample(
    filename: str, content: bytes, limit: int = 3000
) -> tuple[str, list[str]]:
    """返回 (多点采样文本, 标题列表)。"""
    suffix = Path(filename).suffix.lower()
    full_text = ""
    headings: list[str] = []
    try:
        if suffix == ".docx":
            full_text, headings = _docx_full_text(content)
        elif suffix == ".pdf":
            full_text = _pdf_sample(content, limit)
        elif suffix in {".xlsx", ".xlsm", ".xls", ".csv", ".tsv"}:
            full_text = _sheet_preview(filename, content, limit)
    except Exception:
        full_text = ""
    if not full_text:
        full_text = content[: limit * 8].decode("utf-8-sig", errors="replace")
    return _sample_three_points(full_text, limit), headings


def _sample_three_points(text: str, limit: int = 3000) -> str:
    """头部 1500 + 中部 1000 + 尾部 500 字采样。"""
    text = text.strip()
    if len(text) <= limit:
        return text
    head_n = limit // 2          # 1500
    mid_n = limit // 3           # 1000
    tail_n = limit - head_n - mid_n  # 500
    head = text[:head_n]
    mid_start = max(head_n, (len(text) - mid_n) // 2)
    mid = text[mid_start: mid_start + mid_n]
    tail = text[-tail_n:]
    return f"{head}\n……【中部采样】……\n{mid}\n……【尾部采样】……\n{tail}"


def preview_classify_bytes(filename: str, content: bytes) -> dict[str, Any]:
    text = extract_preview_text(filename, content)
    return preview_classify_text(filename, text)


async def preview_classify_with_llm(filename: str, content: bytes, router: object) -> dict[str, Any]:
    """Full classification: keyword pre-screen + LLM (default path)."""
    from backend.classifier import classify_document, classification_to_dict

    text, headings = extract_preview_sample(filename, content)
    result = await classify_document(filename, text, router=router, headings=headings)

    # Also run legacy profile/contract/party detection for extra context
    contract_types, type_evidence, contract_confidence = _classify_profiles(filename, text)
    our_party, party_evidence, party_confidence = _classify_party(f"{filename}\n{text}")
    document_profile = profile_document(filename, text)

    # v1.2：分类有分歧（needs_confirmation）时不再静默自动应用，交给用户确认
    auto_apply_source = (
        not result.needs_confirmation
        and result.confidence >= _SOURCE_AUTO_THRESHOLD
    )
    auto_apply_contract = bool(contract_types) and contract_confidence >= _CONTRACT_AUTO_THRESHOLD
    auto_apply_party = party_confidence >= _PARTY_AUTO_THRESHOLD

    evidence = list(result.evidence)
    evidence.extend(type_evidence[:2])
    if party_evidence:
        evidence.append(party_evidence)

    return {
        "filename": filename,
        "suggested_source_tag": result.source_tag,
        "suggested_contract_types": contract_types,
        "suggested_our_party": our_party,
        "confidence": round(result.confidence, 2),
        "source_confidence": round(result.confidence, 2),
        "contract_confidence": round(contract_confidence, 2),
        "party_confidence": round(party_confidence, 2),
        "auto_apply": auto_apply_source,
        "auto_apply_source": auto_apply_source,
        "auto_apply_contract": auto_apply_contract,
        "auto_apply_party": auto_apply_party,
        "suggested_is_case": result.is_case,
        "suggested_is_redline": result.is_redline,
        "evidence": evidence,
        "document_profile": document_profile,
        # New classification fields
        "classification": classification_to_dict(result),
    }


def preview_classify_text(filename: str, text: str) -> dict[str, Any]:
    """Sync classification: keyword pre-screen only (fallback when no LLM)."""
    from backend.classifier import classify_document_sync, classification_to_dict

    clf = classify_document_sync(filename, text)
    contract_types, type_evidence, contract_confidence = _classify_profiles(filename, text)
    our_party, party_evidence, party_confidence = _classify_party(f"{filename}\n{text}")
    document_profile = profile_document(filename, text)

    auto_apply_source = clf.confidence >= _SOURCE_AUTO_THRESHOLD
    auto_apply_contract = bool(contract_types) and contract_confidence >= _CONTRACT_AUTO_THRESHOLD
    auto_apply_party = party_confidence >= _PARTY_AUTO_THRESHOLD

    evidence = list(clf.evidence)
    evidence.extend(type_evidence[:2])
    if party_evidence:
        evidence.append(party_evidence)

    return {
        "filename": filename,
        "suggested_source_tag": clf.source_tag,
        "suggested_contract_types": contract_types,
        "suggested_our_party": our_party,
        "confidence": round(clf.confidence, 2),
        "source_confidence": round(clf.confidence, 2),
        "contract_confidence": round(contract_confidence, 2),
        "party_confidence": round(party_confidence, 2),
        "auto_apply": auto_apply_source,
        "auto_apply_source": auto_apply_source,
        "auto_apply_contract": auto_apply_contract,
        "auto_apply_party": auto_apply_party,
        "suggested_is_case": clf.is_case,
        "suggested_is_redline": clf.is_redline,
        "evidence": evidence,
        "document_profile": document_profile,
        "classification": classification_to_dict(clf),
    }


def _docx_full_text(content: bytes, max_chars: int = 60_000) -> tuple[str, list[str]]:
    """全文（截断到 max_chars）+ 标题列表，供多点采样。"""
    import docx

    document = docx.Document(io.BytesIO(content))
    parts: list[str] = []
    headings: list[str] = []
    total = 0
    for paragraph in document.paragraphs:
        txt = paragraph.text.strip()
        if not txt:
            continue
        style_name = str(paragraph.style.name) if paragraph.style else ""
        if "Heading" in style_name and len(headings) < 30:
            headings.append(txt[:60])
        parts.append(txt)
        total += len(txt)
        if total >= max_chars:
            break
    if total < 2000:
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
                    total += len(parts[-1])
                if total >= max_chars:
                    break
            if total >= max_chars:
                break
    return "\n".join(parts), headings


def _pdf_sample(content: bytes, limit: int) -> str:
    """PDF 取首页 + 中间页 + 末页文本，供多点采样。"""
    import fitz

    doc = fitz.open(stream=content, filetype="pdf")
    n = len(doc)
    if n == 0:
        return ""
    page_indexes = sorted({0, n // 2, n - 1})
    parts: list[str] = []
    for idx in page_indexes:
        parts.append(doc[idx].get_text("text"))
        if sum(len(p) for p in parts) >= limit * 4:
            break
    return "\n".join(parts)


def _sheet_preview(filename: str, content: bytes, limit: int) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".csv", ".tsv"}:
        return content[: limit * 4].decode("utf-8-sig", errors="replace")[:limit]

    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets[:2]:
        for row in ws.iter_rows(max_row=30, values_only=True):
            text = " ".join(str(cell) for cell in row if cell is not None)
            if text:
                parts.append(text)
            if sum(len(p) for p in parts) >= limit:
                return "\n".join(parts)[:limit]
    return "\n".join(parts)[:limit]


def _classify_source(filename: str, text: str) -> tuple[str, str, float]:
    filename_hits = _source_rule_hits(filename, "filename")
    body_hits = _source_rule_hits(text, "body")

    scores: list[tuple[str, float, list[str]]] = []
    for rule in _source_rules():
        label = str(rule["label"])
        fn_hits = filename_hits.get(label, [])
        txt_hits = body_hits.get(label, [])
        score = min(0.92, len(fn_hits) * 0.58 + len(txt_hits) * 0.16)
        if label == "案例" and _CASE_FILENAME_RX.search(filename):
            score = max(score, 0.72)
            fn_hits = [*fn_hits, "诉...纠纷案"]
        if score > 0:
            hits = [f"文件名:{h}" for h in fn_hits[:3]]
            hits.extend(f"正文:{h}" for h in txt_hits[:2])
            scores.append((label, score, hits))

    if not scores:
        return "历史合同", "未命中强来源关键词，默认历史合同", 0.18

    scores.sort(key=lambda item: item[1], reverse=True)
    label, score, hits = scores[0]
    has_filename_hit = any(hit.startswith("文件名:") for hit in hits)
    if not has_filename_hit and score < _SOURCE_AUTO_THRESHOLD:
        return "历史合同", f"来源线索仅来自正文（{label}: {', '.join(hits)}），默认历史合同", score
    if score < 0.30:
        return "历史合同", f"来源线索较弱（{label}: {', '.join(hits)}），默认历史合同", score
    return label, f"命中{label}: {', '.join(hits)}", score


def _source_rule_hits(text: str, section: str) -> dict[str, list[str]]:
    hits: dict[str, list[str]] = {}
    for rule in _source_rules():
        label = str(rule["label"])
        keywords = rule["filename"] if section == "filename" else rule["body"]
        matched = [kw for kw in keywords if kw and kw in text]
        if matched:
            hits[label] = matched
    return hits


def _classify_profiles(filename: str, text: str) -> tuple[list[str], list[str], float]:
    """v1.2 去噪：泛词不计分；强关键词需 ≥2 个不同词命中才允许自动应用。"""
    scores: list[tuple[str, float, list[str]]] = []
    distinct_strong: dict[str, int] = {}
    for profile in _load_profiles():
        hits: list[str] = []
        raw_score = 0.0
        label = str(profile["label"])
        strong_hits = 0

        for word in _PROFILE_STRONG_KEYWORDS.get(label, ()):
            fn_count = filename.count(word)
            body_count = text.count(word)
            if fn_count or body_count:
                strong_hits += 1
            if fn_count:
                raw_score += min(fn_count, 2) * 8
                hits.append(f"文件名:{word}x{fn_count}")
            if body_count:
                raw_score += min(body_count, 3) * 3
                if len(hits) < 8:
                    hits.append(f"正文:{word}x{body_count}")

        for word in profile["vocabulary"]:
            if not word or word in _GENERIC_PROFILE_TERMS:
                continue  # 泛词（验收/服务/担保等）直接不计分
            fn_count = len(re.findall(re.escape(word), filename, flags=re.IGNORECASE))
            body_count = len(re.findall(re.escape(word), text, flags=re.IGNORECASE))
            if not fn_count and not body_count:
                continue
            raw_score += min(fn_count, 2) * 3 + min(body_count, 3) * 1.2
            count = fn_count + body_count
            if count:
                if len(hits) < 6:
                    hits.append(f"{word}x{count}")
        if label and label in filename:
            raw_score += 12
            strong_hits += 1
            hits.insert(0, f"{label}x1")
        elif label and label in text:
            raw_score += 5
            hits.insert(0, f"{label}x1")
        if raw_score > 0:
            scores.append((label, raw_score, hits))
            distinct_strong[label] = strong_hits

    scores.sort(key=lambda item: item[1], reverse=True)
    if not scores or scores[0][1] < 5:
        evidence = [
            f"弱命中{label}: {', '.join(hits[:4])} (+{score:.1f})"
            for label, score, hits in scores[:3]
        ]
        return [], evidence, 0.0

    top_score = scores[0][1]
    selected = [scores[0][0]]
    if len(scores) > 1 and scores[1][1] >= 0.72 * top_score and scores[1][1] >= 8:
        selected.append(scores[1][0])
    evidence = [
        f"命中{label}: {', '.join(hits[:4])} (+{score:.1f})"
        for label, score, hits in scores[:3]
    ]
    gap = top_score - (scores[1][1] if len(scores) > 1 else 0)
    confidence = min(0.95, 0.38 + top_score / 32)
    if len(scores) > 1 and gap < 3:
        confidence = min(confidence, 0.58)
    # v1.2：强关键词命中不足 2 个不同词 → 置信封顶在自动应用阈值之下（只建议不自动套用）
    if distinct_strong.get(scores[0][0], 0) < 2:
        confidence = min(confidence, _CONTRACT_AUTO_THRESHOLD - 0.05)
    return selected, evidence, confidence


def _classify_party(haystack: str) -> tuple[str, str | None, float]:
    counts: list[tuple[str, int, str]] = []
    for label, keywords in _PARTY_RULES:
        total = sum(haystack.count(keyword) for keyword in keywords)
        if total:
            counts.append((label, total, "/".join(keywords)))
    if not counts:
        return "通用", None, 0.0
    counts.sort(key=lambda item: item[1], reverse=True)
    label, total, keywords = counts[0]
    confidence = min(0.85, 0.45 + total * 0.08)
    return label, f"我方立场线索: {keywords} 命中 {total} 次", confidence


def _load_profiles() -> list[dict[str, Any]]:
    profiles: list[dict[str, Any]] = []
    if not PROFILES_DIR.exists():
        return profiles
    for path in sorted(PROFILES_DIR.glob("*.yaml")):
        try:
            raw = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
        except (OSError, yaml.YAMLError):
            continue
        vocabulary = raw.get("vocabulary", [])
        if isinstance(vocabulary, str):
            vocabulary = [line.strip() for line in vocabulary.splitlines() if line.strip()]
        profiles.append({
            "label": raw.get("name") or path.stem,
            "vocabulary": [str(item).strip() for item in vocabulary if str(item).strip()],
        })
    return profiles
