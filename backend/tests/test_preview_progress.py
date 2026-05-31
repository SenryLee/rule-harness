from __future__ import annotations

from types import SimpleNamespace

from backend.orchestrator import BatchProgress, FidelityStats, PipelineFileState, PipelineState, _pipeline_units
from backend.config import load_config
from backend.parsers import ContentBlock, ParsedDocument, parse_file
from backend.pipelines.p1_body import P1BodyPipeline
from backend.pipelines.p5_case import P5CasePipeline
from backend.preview import preview_classify_text


def test_preview_classify_huarun_manual_as_construction_template():
    result = preview_classify_text(
        "华润置地合同审核指引手册.docx",
        "发包人 总承包 工程量清单 施工总承包 竣工验收 工期索赔 工程变更",
    )

    assert result["suggested_source_tag"] == "标准条款库"
    assert "建工·总包" in result["suggested_contract_types"]
    assert result["suggested_our_party"] == "发包人"
    assert result["confidence"] >= 0.6
    assert result["auto_apply_source"] is True
    assert result["auto_apply_contract"] is True
    assert result["evidence"]


def test_preview_classify_redline_file_source_tag():
    result = preview_classify_text(
        "谈判红线底线.txt",
        "客户要求无限责任时不可接受，可退让至赔偿上限。",
    )

    assert result["suggested_source_tag"] == "公司红线"


def test_preview_classify_case_file_source_tag():
    result = preview_classify_text(
        "买卖合同纠纷判决案例.txt",
        "法院认为该违约金过高，依法予以调整。",
    )

    assert result["suggested_source_tag"] == "案例"


def test_preview_classify_gift_property_not_as_construction():
    result = preview_classify_text(
        "夫妻间赠与房产的处理.docx",
        "夫妻一方将名下房产赠与另一方，涉及不动产登记、赠与撤销与共有财产认定。",
    )

    assert result["suggested_source_tag"] == "历史合同"
    assert result["suggested_contract_types"][0] == "房地产"
    assert "建工·总包" not in result["suggested_contract_types"]


def test_preview_classify_case_analysis_not_as_construction():
    result = preview_classify_text(
        "附解除条件赠与合同的认定边界与证据适用——基于一起婚姻关联动产赠与纠纷案的分析.docx",
        "法院认为附解除条件赠与合同应结合证据规则、履行行为与合同无效抗辩进行审查。",
    )

    assert result["suggested_source_tag"] == "案例"
    assert result["suggested_contract_types"][0] == "通用商事"
    assert "建工·总包" not in result["suggested_contract_types"]


def test_preview_classify_bylaw_file_not_as_case():
    result = preview_classify_text(
        "赠与公证细则.docx",
        "公证机构办理赠与合同公证，应当核验当事人身份、赠与意思表示与财产权属证明。",
    )

    assert result["suggested_source_tag"] == "法规"
    assert result["auto_apply_source"] is True


def test_preview_low_confidence_suggestion_is_not_auto_applied():
    result = preview_classify_text(
        "材料.docx",
        "争议 变更 协商",
    )

    assert result["auto_apply_contract"] is False


def test_batch_progress_exposes_default_pipeline_rows():
    progress = BatchProgress(total_files=2)
    payload = progress.to_dict()

    assert list(payload["pipeline_progress"].keys()) == [
        "P1",
        "P2",
        "P3",
        "P4",
        "P5",
        "direct",
    ]
    assert payload["pipeline_progress"]["P1"]["label"] == "正文抽取"
    assert payload["fidelity_stats"] == {
        "intercepted": 0,
        "placeholders": 0,
        "discarded": 0,
        "voice_mismatch": 0,
    }


def test_pipeline_state_serializes_file_cells():
    state = PipelineState(label="正文抽取")
    progress = BatchProgress(
        pipeline_progress={"P1": state},
        fidelity_stats=FidelityStats(discarded=2),
    )
    progress.pipeline_progress["P1"].files["a.docx"] = PipelineFileState(
        filename="a.docx",
        status="done",
        blocks_total=3,
        blocks_done=3,
        rules_emitted=7,
    )

    payload = progress.to_dict()
    assert payload["pipeline_progress"]["P1"]["files"]["a.docx"]["rules_emitted"] == 7
    assert payload["fidelity_stats"]["discarded"] == 2


def test_pipeline_block_progress_updates_incrementally():
    progress = BatchProgress(total_blocks=3)
    state = progress.pipeline_progress["P1"]
    state.files_total = 1
    state.blocks_total = 3
    state.files["a.docx"] = PipelineFileState(filename="a.docx", status="running", blocks_total=3)

    progress.mark_pipeline_block_done("P1", "a.docx", rules_emitted=2)
    progress.mark_pipeline_block_done("P1", "a.docx", rules_emitted=1)
    progress.mark_pipeline_done("P1", "a.docx", rules_emitted=3)

    payload = progress.to_dict()["pipeline_progress"]["P1"]
    assert payload["blocks_done"] == 3
    assert payload["rules_emitted"] == 3
    assert progress.processed_blocks == 3


def test_batch_progress_accumulates_token_usage():
    progress = BatchProgress()

    progress.add_token_usage({"total_tokens": 12})
    progress.add_token_usage({"prompt_tokens": 5, "completion_tokens": 7})
    progress.add_token_usage({})

    assert progress.tokens_used == 24


def test_p1_llm_failure_is_reported_to_batch_progress():
    class FakePrimary:
        name = "fake"

    class FailingRouter:
        primary = FakePrimary()

        async def chat_json(self, system, user, temperature):
            raise RuntimeError("HTTP Error 402: Payment Required")

    doc = ParsedDocument(
        sha256="x",
        filename="合同.docx",
        source_tag="历史合同",
        priority=5,
        contract_types=["通用商事"],
        industry_context=None,
        is_scanned=False,
        blocks=(ContentBlock("p1", "付款期限应在验收后30日内完成。", "1", "paragraph"),),
        comments=(),
        revisions=(),
        is_redline_doc=False,
        is_case_doc=False,
        is_passthrough=False,
    )
    progress = BatchProgress()
    pipe = P1BodyPipeline(FailingRouter(), load_config())

    rules = __import__("asyncio").run(
        pipe._extract_block(doc, doc.blocks[0], {"progress": progress})
    )

    assert rules == []
    assert progress.errors
    assert progress.errors[0].startswith("llm_failed:P1:合同.docx:p1:RuntimeError")
    assert "402" in progress.errors[0]


def test_parse_legacy_doc_uses_available_text_converter(tmp_path, monkeypatch):
    path = tmp_path / "旧版合同.doc"
    path.write_bytes(b"legacy-doc-binary")

    monkeypatch.setattr(
        "backend.parsers.shutil.which",
        lambda tool: "/usr/bin/textutil" if tool == "textutil" else None,
    )

    def fake_run(cmd, check, capture_output, text, timeout):
        assert cmd[:4] == ["textutil", "-convert", "txt", "-stdout"]
        return SimpleNamespace(returncode=0, stdout="付款条款\n\n违约责任", stderr="")

    monkeypatch.setattr("backend.parsers.subprocess.run", fake_run)

    parsed = parse_file(path, source_tag="历史合同", contract_types=["通用商事"])

    assert parsed.parse_warnings == ()
    assert [block.text for block in parsed.blocks] == ["付款条款", "违约责任"]


def test_parse_docx_keeps_table_rows_for_extraction(tmp_path):
    import docx

    path = tmp_path / "合同条款.docx"
    document = docx.Document()
    document.add_paragraph("付款安排")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "条款"
    table.cell(0, 1).text = "内容"
    table.cell(1, 0).text = "违约金"
    table.cell(1, 1).text = "买方逾期付款的，应当按日支付违约金。"
    document.save(path)

    parsed = parse_file(path, source_tag="历史合同", contract_types=["通用商事"])

    assert parsed.is_passthrough is False
    assert any(block.block_type == "table_row" for block in parsed.blocks)
    assert any("买方逾期付款" in block.text for block in parsed.blocks)


def test_parse_docx_rule_table_as_passthrough(tmp_path):
    import docx

    path = tmp_path / "审查清单.docx"
    document = docx.Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "风险等级"
    table.cell(0, 1).text = "检查项"
    table.cell(0, 2).text = "审查要求"
    table.cell(1, 0).text = "高"
    table.cell(1, 1).text = "违约金上限"
    table.cell(1, 2).text = "违约金总额不得超过合同总价的 30%。"
    document.save(path)

    parsed = parse_file(path, source_tag="标准条款库", contract_types=["通用商事"])

    assert parsed.is_passthrough is True
    assert parsed.blocks[0].block_type == "table_row"
    assert "风险等级: 高" in parsed.blocks[0].text


def test_p1_also_processes_case_docs_for_coverage():
    doc = ParsedDocument(
        sha256="x",
        filename="案例分析.docx",
        source_tag="案例",
        priority=5,
        contract_types=["通用商事"],
        industry_context=None,
        is_scanned=False,
        blocks=(),
        comments=(),
        revisions=(),
        is_redline_doc=False,
        is_case_doc=True,
        is_passthrough=False,
    )
    pipe = P1BodyPipeline.__new__(P1BodyPipeline)

    assert pipe.applicable(doc)


def test_p5_processes_case_in_chunks_with_local_anchors():
    import asyncio

    class FakePrimary:
        name = "fake-model"

    class FakeRouter:
        primary = FakePrimary()

        def __init__(self):
            self.calls = 0

        async def chat_json(self, system, user, temperature):
            self.calls += 1
            return {
                "rules": [
                    {
                        "risk_level": "中",
                        "keywords": ["赠与", "证据", "登记"],
                        "check_item": f"证据链是否完整{self.calls}",
                        "requirement": "[合规] 审查赠与事实时应核验证据链",
                        "notes": "根据本分块裁判要点提炼，不编造案号。",
                        "rule_type": "governance",
                        "theme_key": "delivery.title_transfer.point",
                        "subject": f"审查人员{self.calls}",
                        "predicate": "应核验",
                        "threshold_type": "无",
                        "direction": "正向",
                        "cited_cases": [],
                        "self_confidence": 0.65,
                        "uncertainty_points": [],
                    }
                ]
            }

    doc = ParsedDocument(
        sha256="x",
        filename="案例分析.docx",
        source_tag="案例",
        priority=5,
        contract_types=["通用商事"],
        industry_context=None,
        is_scanned=False,
        blocks=tuple(
            ContentBlock(
                block_id=f"p{i}",
                text=f"第{i}段 法院认为赠与合同证据链应结合登记、交付和履行事实综合判断。" * 8,
                location=str(i),
                block_type="paragraph",
            )
            for i in range(8)
        ),
        comments=(),
        revisions=(),
        is_redline_doc=False,
        is_case_doc=True,
        is_passthrough=False,
    )
    router = FakeRouter()
    pipe = P5CasePipeline(router, load_config())

    rules = asyncio.run(pipe.extract(doc, {}))

    assert router.calls >= 2
    assert len(rules) == router.calls
    assert _pipeline_units("P5", doc) == router.calls
    assert len({rule.source_location for rule in rules}) == router.calls
    assert any("第7段" in rule.source_excerpt for rule in rules)
    assert all(rule.rule_type == "governance" for rule in rules)
